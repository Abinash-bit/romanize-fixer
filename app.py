import streamlit as st
import re
import copy
import io
import os
import time
import logging
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import Counter

from suggest import SuggestEngine
import romanize_srt

# Load ANTHROPIC_API_KEY (and friends) from a local .env if present. python-dotenv
# strips the whitespace around `KEY = value`, so the existing .env format works.
try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:  # pragma: no cover - optional dependency
    pass

# ─────────────────────────────────────────────
# Logging — shows in the terminal where you ran `streamlit run app.py`
# ─────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-5s | %(name)s | %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("romanize.app")

# ─────────────────────────────────────────────
# Tier styling
# ─────────────────────────────────────────────
# exact   : confirmed dictionary correction        -> yellow highlight
# fix     : engine fix, higher confidence (review) -> green  highlight
# suggest : engine guess, lower confidence (review)-> cyan   highlight
TIER_HL = {"exact": "yellow", "fix": "green", "suggest": "cyan"}

FIX_CUTOFF = 0.75       # >= this -> "fix" tier (green). below -> "suggest"
SUGGEST_FLOOR = 0.60    # below this -> ignore entirely (the 0.50-0.60 band was
                        # ~50% accurate — a coin flip — so we don't show it)


# ─────────────────────────────────────────────
# Dictionary loading
# ─────────────────────────────────────────────
def _clean_note(s: str) -> str:
    """Strip human annotation notes like [accent] / (as per dialogue)."""
    s = re.sub(r"\[[^\]]*\]", "", s)
    s = re.sub(r"\([^)]*\)", "", s)
    return s.strip()


def load_dictionary(file_bytes: bytes):
    """Return (exact_replacements, pairs).

    exact_replacements: {variant_lower: planetread_word}  (verbatim, as before)
    pairs:              [(wrong, right), ...]  cleaned, for the learning engine
    """
    t0 = time.perf_counter()
    log.info("Loading dictionary (%d KB)…", len(file_bytes) // 1024)
    doc = Document(io.BytesIO(file_bytes))
    replacements = {}
    pairs = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 3:
                continue
            ai_word, pr_word = cells[1].strip(), cells[2].strip()
            if not ai_word or not pr_word:
                continue
            if ai_word.lower() in ("ai version", "ai_version"):
                continue
            for variant in ai_word.split('/'):
                variant = variant.strip()
                if variant:
                    replacements[variant.lower()] = pr_word
            a, b = _clean_note(ai_word), _clean_note(pr_word)
            if a and b:
                pairs.append((a, b))
    log.info("Dictionary loaded: %d exact variants, %d clean pairs in %.2fs",
             len(replacements), len(pairs), time.perf_counter() - t0)
    return replacements, pairs


# ─────────────────────────────────────────────
# Run XML builders
# ─────────────────────────────────────────────
def _base_rPr(source_run_xml=None):
    if source_run_xml is not None:
        src = source_run_xml.find(qn('w:rPr'))
        if src is not None:
            return copy.deepcopy(src)
    return OxmlElement('w:rPr')


def _make_t(text: str):
    t = OxmlElement('w:t')
    t.text = text
    if text.startswith(' ') or text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return t


def _clean_rPr(rPr):
    for tag in ('w:color', 'w:strike', 'w:highlight'):
        for el in rPr.findall(qn(tag)):
            rPr.remove(el)
    return rPr


def build_red_run(text, source_run_xml=None):
    r = OxmlElement('w:r')
    rPr = _clean_rPr(_base_rPr(source_run_xml))
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'FF0000')
    rPr.append(color)
    r.append(rPr)
    r.append(_make_t(text))
    return r


def build_highlight_run(text, color, source_run_xml=None):
    r = OxmlElement('w:r')
    rPr = _clean_rPr(_base_rPr(source_run_xml))
    hl = OxmlElement('w:highlight')
    hl.set(qn('w:val'), color)
    rPr.append(hl)
    r.append(rPr)
    r.append(_make_t(text))
    return r


def build_normal_run(text, source_run_xml=None):
    r = OxmlElement('w:r')
    r.append(_base_rPr(source_run_xml))
    r.append(_make_t(text))
    return r


# ─────────────────────────────────────────────
# Per-paragraph replacement
# ─────────────────────────────────────────────
def build_exact_regex(replacements):
    """Compile ONE regex matching any dictionary key (longest first). Built once
    per run instead of 600 separate regexes per paragraph — the big speed win."""
    keys = sorted(replacements.keys(), key=len, reverse=True)
    if not keys:
        return None
    alt = "|".join(re.escape(k) for k in keys)
    return re.compile(r'(?<![A-Za-z])(' + alt + r')(?![A-Za-z])', re.IGNORECASE)


def replace_in_paragraph(para, exact_re, replacements, engine, enable_engine):
    """Return list of change dicts: {orig, repl, tier, conf, reason}."""
    full_text = "".join(run.text for run in para.runs)
    if not full_text.strip():
        return []

    # 1) exact dictionary spans — single regex pass, matches are non-overlapping
    spans = []  # (start, end, replacement, tier, conf, reason)
    if exact_re is not None:
        for m in exact_re.finditer(full_text):
            rep = replacements.get(m.group(1).lower())
            if rep is not None:
                spans.append((m.start(), m.end(), rep, "exact", 1.0, "dictionary"))

    # 2) engine suggestions on remaining words
    if enable_engine and engine is not None:
        for m in re.finditer(r'[A-Za-z]+', full_text):
            if any(s < m.end() and m.start() < e for s, e, *_ in spans):
                continue
            sug = engine.suggest(m.group(0))
            if sug is None or sug.confidence < SUGGEST_FLOOR:
                continue
            tier = "fix" if sug.confidence >= FIX_CUTOFF else "suggest"
            spans.append((m.start(), m.end(), sug.correction, tier, sug.confidence, sug.reason))

    if not spans:
        return []

    spans.sort(key=lambda x: x[0])
    changes = [
        {"orig": full_text[s:e], "repl": rep, "tier": tier, "conf": conf, "reason": reason}
        for s, e, rep, tier, conf, reason in spans
    ]

    # rebuild the paragraph
    first_run_xml = para.runs[0]._r if para.runs else None
    p_xml = para._p
    for r in p_xml.findall(qn('w:r')):
        p_xml.remove(r)

    pos = 0
    for (start, end, replacement, tier, conf, reason) in spans:
        if pos < start:
            p_xml.append(build_normal_run(full_text[pos:start], first_run_xml))
        p_xml.append(build_red_run(full_text[start:end], first_run_xml))
        p_xml.append(build_normal_run(' ', first_run_xml))
        p_xml.append(build_highlight_run(replacement, TIER_HL[tier], first_run_xml))
        pos = end
    if pos < len(full_text):
        p_xml.append(build_normal_run(full_text[pos:], first_run_xml))

    return changes


def process_document(dict_bytes, input_bytes, enable_engine, progress_cb=None):
    """progress_cb(done, total, stage_msg) is called as work proceeds (optional)."""
    t_start = time.perf_counter()

    def report(done, total, msg):
        if progress_cb:
            progress_cb(done, total, msg)

    report(0, 1, "Reading dictionary…")
    replacements, pairs = load_dictionary(dict_bytes)
    exact_re = build_exact_regex(replacements)

    report(0, 1, "Training pattern engine…")
    engine = SuggestEngine(pairs) if enable_engine else None

    report(0, 1, "Opening subtitle file…")
    doc = Document(io.BytesIO(input_bytes))

    # gather every paragraph up-front so we know the total for the progress bar
    paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    total = len(paras)
    log.info("Scanning %d paragraphs (engine=%s)…", total, "on" if enable_engine else "off")

    all_changes = []
    t_scan = time.perf_counter()
    for i, para in enumerate(paras, 1):
        all_changes.extend(replace_in_paragraph(para, exact_re, replacements, engine, enable_engine))
        if i % 200 == 0 or i == total:
            log.info("  …%d/%d paragraphs, %d changes so far", i, total, len(all_changes))
            report(i, total, f"Scanning paragraphs… {i}/{total}")
    log.info("Scan done: %d changes in %.2fs", len(all_changes), time.perf_counter( ) - t_scan)

    if engine is not None:
        log.info("Engine cache: %d unique words looked up", len(engine._cache))

    report(total, total, "Saving fixed file…")
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    log.info("process_document finished in %.2fs total", time.perf_counter() - t_start)
    return out, replacements, all_changes


def build_dictionary_export(changes):
    """Make a 3-column docx (same format) of engine-found pairs to paste back."""
    seen, rows = set(), []
    for ch in changes:
        if ch["tier"] == "exact":
            continue
        key = (ch["orig"].lower(), ch["repl"])
        if key in seen:
            continue
        seen.add(key)
        rows.append((ch["orig"], ch["repl"]))

    doc = Document()
    doc.add_heading("Engine-found corrections (review before merging)", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "", "AI Version", "Planet Read Version"
    for i, (a, b) in enumerate(rows, 1):
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = str(i), a, b
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out, len(rows)


# regex for an SRT timecode line: 00:00:01,200 --> 00:00:04,000
TIMECODE_RE = re.compile(
    r'\d{1,2}:\d{2}:\d{2}[,.]\d{3}\s*-->\s*\d{1,2}:\d{2}:\d{2}[,.]\d{3}'
)


def _match_case(matched, replacement):
    """Make a correction follow the case of the word it replaces, so it is
    capitalised only where the original word was (start of a sentence / proper
    noun) and stays lowercase mid-sentence — instead of the dictionary's stored
    capitalisation leaking capital letters into the middle of a sentence."""
    if not replacement:
        return replacement
    rep = replacement.lower()                 # canonical lowercase base form
    if matched[:1].isupper():
        rep = rep[:1].upper() + rep[1:]        # original was capitalised -> match it
    return rep


def build_srt(input_bytes, replacements):
    """Build a clean .srt from the subtitle Word file.

    - applies ONLY the confirmed dictionary (yellow) corrections — wrong word is
      replaced by the correct word, with NO red/green/cyan markup at all
    - engine guesses are NOT applied (original word kept) — they're unconfirmed
    - timecodes / block numbers already present in the Word text are preserved

    Returns (srt_bytes, num_timecodes, num_corrections).
    """
    exact_re = build_exact_regex(replacements)
    n_fixes = 0

    def fix_line(text):
        nonlocal n_fixes
        if exact_re is None:
            return text

        def repl(m):
            nonlocal n_fixes
            rep = replacements.get(m.group(1).lower())
            if rep is None:
                return m.group(0)
            n_fixes += 1
            return _match_case(m.group(1), rep)
        return exact_re.sub(repl, text)

    doc = Document(io.BytesIO(input_bytes))
    lines = [fix_line(p.text) for p in doc.paragraphs]
    srt_text = "\n".join(lines).strip() + "\n"
    # tidy whitespace: collapse runs of spaces between words, trim per-line edges
    srt_text = re.sub(r'[ \t]{2,}', ' ', srt_text)
    srt_text = re.sub(r'(?m)^[ \t]+|[ \t]+$', '', srt_text)
    n_codes = len(TIMECODE_RE.findall(srt_text))
    log.info("SRT built: %d timecodes, %d dictionary corrections applied", n_codes, n_fixes)
    return srt_text.encode("utf-8"), n_codes, n_fixes


# ─────────────────────────────────────────────
# Streamlit UI
# ─────────────────────────────────────────────
st.set_page_config(page_title="PlanetRead – Romanizer", page_icon="📖", layout="wide")


def resolve_api_key():
    """Find the Anthropic key from the environment (.env locally) or from Streamlit
    Secrets (the cloud path), and export it so anthropic.Anthropic() picks it up."""
    key = os.getenv("ANTHROPIC_API_KEY")
    if not key:
        try:
            key = st.secrets.get("ANTHROPIC_API_KEY")
        except Exception:
            key = None
        if key:
            os.environ["ANTHROPIC_API_KEY"] = key
    return key


def _render_readonly_tier(title, items, color, note):
    """Exact dictionary fixes — always applied, shown for reference (no checkbox)."""
    if not items:
        return
    st.markdown(f"#### {title}")
    if note:
        st.caption(note)
    counts = Counter((c["orig"], c["repl"]) for c in items)
    for (orig, corr), cnt in sorted(counts.items(), key=lambda x: -x[1]):
        c1, c2, c3 = st.columns([3, 3, 1])
        c1.markdown(f"<span style='color:red;font-weight:bold'>{orig}</span>", unsafe_allow_html=True)
        c2.markdown(f"<span style='background:{color};padding:2px 10px;border-radius:4px;font-weight:bold'>{corr}</span>", unsafe_allow_html=True)
        c3.markdown(f"×**{cnt}**")


def _render_checkbox_tier(title, items, color, tier, scan_id):
    """Render an engine tier with a tick column. Returns the accepted (orig, repl) set.

    Keys include scan_id so a fresh scan starts every box unticked instead of
    inheriting stale state mapped to different rows.
    """
    sel = set()
    if not items:
        return sel
    st.markdown(f"#### {title}")
    counts = Counter((c["orig"], c["repl"]) for c in items)
    confs = {(c["orig"], c["repl"]): c["conf"] for c in items}
    rows = sorted(counts.items(), key=lambda x: -x[1])

    all_on = st.checkbox(f"Accept all {len(rows)}", key=f"selall_{scan_id}_{tier}")

    h0, h1, h2, h3, h4 = st.columns([0.7, 3, 3, 1, 1])
    h0.markdown("<b>Use</b>", unsafe_allow_html=True)
    h1.markdown("<b>Wrong (in file)</b>", unsafe_allow_html=True)
    h2.markdown("<b>Correction</b>", unsafe_allow_html=True)
    h3.markdown("<b>Conf.</b>", unsafe_allow_html=True)
    h4.markdown("<b>Count</b>", unsafe_allow_html=True)
    for i, ((orig, corr), cnt) in enumerate(rows):
        c0, c1, c2, c3, c4 = st.columns([0.7, 3, 3, 1, 1])
        checked = c0.checkbox("keep", key=f"sel_{scan_id}_{tier}_{i}",
                              disabled=all_on, label_visibility="collapsed")
        c1.markdown(f"<span style='color:red;font-weight:bold'>{orig}</span>", unsafe_allow_html=True)
        c2.markdown(f"<span style='background:{color};padding:2px 10px;border-radius:4px;font-weight:bold'>{corr}</span>", unsafe_allow_html=True)
        c3.markdown(f"{confs[(orig, corr)]*100:.0f}%")
        c4.markdown(f"**{cnt}**")
        if all_on or checked:
            sel.add((orig, corr))
    return sel


def render_selection(all_changes, n_pairs, scan_id):
    """Metrics + always-applied exact tier + tickable engine tiers. Returns the
    set of accepted engine (orig, repl) pairs."""
    by_tier = {"exact": [], "fix": [], "suggest": []}
    for ch in all_changes:
        by_tier[ch["tier"]].append(ch)

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Dictionary Pairs", n_pairs)
    m2.metric("✅ Dictionary Fixes", len(by_tier["exact"]))
    m3.metric("🟢 Engine Fixes", len(by_tier["fix"]))
    m4.metric("🔵 Engine Guesses", len(by_tier["suggest"]))

    _render_readonly_tier("✅ Confirmed Dictionary Fixes (always applied)",
                          by_tier["exact"], "yellow",
                          "Exact matches from your dictionary — always applied.")
    sel = set()
    sel |= _render_checkbox_tier("🟢 Engine Fixes — tick to keep", by_tier["fix"],
                                 "#9bffb0", "fix", scan_id)
    sel |= _render_checkbox_tier("🔵 Engine Guesses — tick to keep", by_tier["suggest"],
                                 "#9bf6ff", "suggest", scan_id)
    if not by_tier["fix"] and not by_tier["suggest"]:
        st.caption("No engine suggestions to review (engine off, or none found).")
    return sel


def build_marked_docx(text, apply_re, apply_map, tier_map):
    """Word review file: each applied correction shown as red original + a
    tier-coloured highlight of the replacement. Only keys in apply_map are marked."""
    doc = Document()
    for line in text.split("\n"):
        para = doc.add_paragraph()
        if not line:
            continue
        spans = []
        if apply_re is not None:
            for m in apply_re.finditer(line):
                key = m.group(1).lower()
                rep = apply_map.get(key)
                if rep is not None:
                    rep = _match_case(m.group(1), rep)
                    spans.append((m.start(), m.end(), rep, tier_map.get(key, "exact")))
        if not spans:
            para._p.append(build_normal_run(line))
            continue
        pos = 0
        for s, e, rep, tier in spans:
            if pos < s:
                para._p.append(build_normal_run(line[pos:s]))
            para._p.append(build_red_run(line[s:e]))
            para._p.append(build_normal_run(' '))
            para._p.append(build_highlight_run(rep, TIER_HL.get(tier, "yellow")))
            pos = e
        if pos < len(line):
            para._p.append(build_normal_run(line[pos:]))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_selected_outputs(romanized_text, replacements, all_changes, selected):
    """Apply the exact dictionary + only the ticked engine pairs. Returns a dict
    with the review docx, final SRT, engine-pairs export, and counts."""
    selected_engine = [
        c for c in all_changes
        if c["tier"] in ("fix", "suggest") and (c["orig"], c["repl"]) in selected
    ]
    # apply_map drives both the markup and the SRT replacement; exact wins conflicts.
    apply_map, tier_map = {}, {}
    for c in selected_engine:
        k = c["orig"].lower()
        apply_map[k] = c["repl"]
        tier_map[k] = c["tier"]
    for k, v in replacements.items():
        apply_map[k] = v
        tier_map[k] = "exact"

    apply_re = build_exact_regex(apply_map)
    review_docx = build_marked_docx(romanized_text, apply_re, apply_map, tier_map)

    rom_docx = romanize_srt.srt_text_to_docx_bytes(romanized_text)
    srt_final, n_codes, n_fixes = build_srt(rom_docx, apply_map)

    exp_bytes, n_new = build_dictionary_export(selected_engine)
    return {
        "review_docx": review_docx,
        "srt_final": srt_final,
        "n_codes": n_codes,
        "n_fixes": n_fixes,
        "engine_pairs": exp_bytes.getvalue(),
        "n_new": n_new,
        "n_engine_applied": len({(c["orig"], c["repl"]) for c in selected_engine}),
    }


def render_page(enable_engine):
    """Single staged flow: romanize the SRT first, then apply the dictionary."""
    st.markdown(
        "Romanize a native-script subtitle (Hindi, Punjabi, Marathi, Tamil, Telugu, "
        "Kannada, Gujarati) with Claude, then apply the dictionary to get the final "
        "corrected `.srt`."
    )

    has_key = bool(resolve_api_key())
    if not romanize_srt.anthropic_available():
        st.error("The `anthropic` package isn't installed. Run "
                 "`pip install -r requirements.txt`.")
    elif not has_key:
        st.warning("⚠️ No `ANTHROPIC_API_KEY` found. Add it to a `.env` file (local) or "
                   "to **Secrets** (Streamlit Cloud), e.g. `ANTHROPIC_API_KEY = sk-ant-...`.")

    # ── ① Upload subtitle + pick language → romanize ─────────────────────
    st.subheader("① Upload subtitle & romanize")
    c1, c2 = st.columns([3, 2])
    with c1:
        srt_file = st.file_uploader("Native-script subtitle (.srt)", type=["srt"], key="srt_input")
    with c2:
        language = st.selectbox("Source language", list(romanize_srt.LANGUAGES.keys()), key="srt_lang")

    can_romanize = bool(srt_file) and has_key and romanize_srt.anthropic_available()
    if srt_file and st.button("🌐 Romanize SRT", type="primary",
                              use_container_width=True, key="run_romanize",
                              disabled=not can_romanize):
        progress = st.progress(0.0, text="Contacting Anthropic…")
        status = st.empty()
        t_ui = time.perf_counter()

        def on_progress(done, total):
            frac = min(1.0, done / total) if total else 0.0
            progress.progress(frac, text=f"Romanizing lines… {done}/{total}")
            status.caption(f"⏱️ {time.perf_counter() - t_ui:.1f}s")

        try:
            srt_bytes = srt_file.read()
            romanized_text, stats = romanize_srt.romanize_srt_bytes(
                srt_bytes, language, progress_cb=on_progress,
            )
            progress.progress(1.0, text="Romanization done!")
            # A fresh romanization invalidates any earlier scan/build.
            st.session_state["srt_result"] = {
                "base": srt_file.name.rsplit(".srt", 1)[0],
                "romanized_text": romanized_text,
                "stats": stats,
            }
        except Exception as e:
            log.exception("SRT romanization failed")
            st.session_state.pop("srt_result", None)
            st.error(f"Something went wrong: {e}")
            st.exception(e)

    if not srt_file:
        st.info("👆 Upload a native-script `.srt` and choose its language.")

    res = st.session_state.get("srt_result")
    if not res:
        return

    # ── show romanized output ────────────────────────────────────────────
    stats = res["stats"]
    msg = f"✅ Romanized {stats['lines']} line(s) across {stats['blocks']} block(s)."
    if stats.get("missing"):
        msg += f" ⚠️ {stats['missing']} line(s) kept as-is (model didn't return them)."
    if stats.get("cached"):
        msg += " ♻️ Loaded from cache — no Anthropic call."
    st.success(msg)
    st.download_button(
        f"⬇️ {res['base']}_romanized.srt", data=res["romanized_text"].encode("utf-8"),
        file_name=f"{res['base']}_romanized.srt", mime="application/x-subrip",
        use_container_width=True, key="dl_srt_romanized",
    )

    # ── ② Upload dictionary → scan for corrections ───────────────────────
    st.divider()
    st.subheader("② Apply the dictionary")
    dict_file = st.file_uploader("Universal Dictionary (.docx)", type=["docx"], key="srt_dict")

    if dict_file and st.button("🔍 Find corrections", type="primary",
                               use_container_width=True, key="run_dict"):
        try:
            with st.spinner("Scanning for dictionary + engine corrections…"):
                dict_bytes = dict_file.read()
                rom_docx = romanize_srt.srt_text_to_docx_bytes(res["romanized_text"])
                # We only need the candidate list + the exact dictionary here; the
                # outputs are built later from the user's ticked selection.
                _doc, replacements, all_changes = process_document(
                    dict_bytes, rom_docx, enable_engine,
                )
            scan_id = st.session_state.get("_scan_seq", 0) + 1
            st.session_state["_scan_seq"] = scan_id
            res.update({
                "replacements": replacements,
                "all_changes": all_changes,
                "n_pairs": len(replacements),
                "scan_id": scan_id,
                "scanned": True,
            })
            res.pop("built", None)   # a new scan invalidates any previously built files
            st.session_state["srt_result"] = res
        except Exception as e:
            log.exception("Dictionary scan failed")
            st.error(f"Something went wrong: {e}")
            st.exception(e)

    if not dict_file:
        st.info("👆 Upload the Universal Dictionary (.docx), then click **Find corrections**.")

    if not res.get("scanned"):
        return

    # ── review: tick which engine suggestions to keep ────────────────────
    st.divider()
    st.markdown("### 🔤 Review corrections")
    st.caption("Yellow dictionary fixes are always applied. Tick the engine "
               "fixes/guesses you want — only ticked ones go into the files below.")
    selected = render_selection(res["all_changes"], res["n_pairs"], res["scan_id"])

    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("📦 Build files with the ticked corrections", type="primary",
                 use_container_width=True, key="build_outputs"):
        with st.spinner("Building files…"):
            res["built"] = build_selected_outputs(
                res["romanized_text"], res["replacements"], res["all_changes"], selected,
            )
        st.session_state["srt_result"] = res

    built = res.get("built")
    if not built:
        st.info("Tick the corrections you want, then click **Build files** above. "
                "(Re-build after changing any tick.)")
        return

    # ── ③ download the built files ───────────────────────────────────────
    base = res["base"]
    d1, d2, d3 = st.columns(3)
    with d1:
        st.markdown("#### 🎞️ Final SRT")
        st.download_button(
            f"⬇️ {base}_romanized_fixed.srt", data=built["srt_final"],
            file_name=f"{base}_romanized_fixed.srt", mime="application/x-subrip",
            use_container_width=True, type="primary", key="dl_srt_final",
        )
        st.caption(f"Romanized + {built['n_fixes']} applied correction(s) "
                   f"({built['n_engine_applied']} ticked engine), "
                   f"{built['n_codes']} timecodes kept.")
    with d2:
        st.markdown("#### 📥 Review Word File")
        st.download_button(
            f"⬇️ {base}_review.docx", data=built["review_docx"],
            file_name=f"{base}_review.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, key="dl_srt_review",
        )
        st.caption("Shows only the applied corrections (yellow + the ticked green/cyan).")
    with d3:
        st.markdown("#### 🔁 Grow the Dictionary")
        st.download_button(
            f"⬇️ Ticked engine pairs ({built['n_new']})",
            data=built["engine_pairs"], file_name="engine_found_pairs.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, disabled=(built["n_new"] == 0), key="dl_srt_pairs",
        )
        st.caption("Only the engine pairs you ticked. Paste into the master dictionary.")


st.markdown("""
    <h1 style='color:#E05C00;'>📖 PlanetRead – Romanizer</h1>
    <p style='font-size:14px;'>
      <span style='background:yellow;padding:2px 8px;border-radius:4px;'>yellow</span>
      = confirmed dictionary fix &nbsp;·&nbsp;
      <span style='background:#9bffb0;padding:2px 8px;border-radius:4px;'>green</span>
      = engine fix (please confirm) &nbsp;·&nbsp;
      <span style='background:#9bf6ff;padding:2px 8px;border-radius:4px;'>cyan</span>
      = engine guess (review carefully)
    </p>
    <hr>
""", unsafe_allow_html=True)

enable_engine = st.sidebar.checkbox("🧠 Enable AI pattern engine", value=True)
st.sidebar.caption(
    "Learns spelling-correction patterns from your dictionary and applies them to "
    "words **not** in the table. Offline, no internet. ~75–80% precise on novel words, "
    "so its fixes are flagged for your review, never silently applied."
)

st.sidebar.divider()
st.sidebar.caption(
    "Romanized SRTs are cached on disk by file + language, so re-running the same "
    "file never calls Anthropic again."
)
if st.sidebar.button("🧹 Clear romanization cache"):
    n = romanize_srt.clear_cache()
    st.session_state.pop("srt_result", None)
    st.sidebar.success(f"Cleared {n} cached romanization(s).")

render_page(enable_engine)

st.markdown("<br><hr><p style='text-align:center; color:#aaa; font-size:13px;'>PlanetRead · Romanized Shabdkosh Tool</p>", unsafe_allow_html=True)
