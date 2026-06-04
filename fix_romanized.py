"""
fix_romanized.py
----------------
Reads a universal dictionary (Roman_Shabdkosh) and a romanized subtitle docx.
For every AI-version word found it:
  1. Keeps the original word in RED + STRIKETHROUGH
  2. Appends the corrected PlanetRead word in YELLOW HIGHLIGHT

Usage:
    python fix_romanized.py <dictionary.docx> <input.docx> <output.docx>
"""

import sys
import re
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ──────────────────────────────────────────────────────────────
# 1. Parse the dictionary
# ──────────────────────────────────────────────────────────────

def load_dictionary(dict_path: str) -> dict:
    doc = Document(dict_path)
    replacements = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 3:
                continue
            ai_word = cells[1].strip()
            pr_word = cells[2].strip()
            if not ai_word or not pr_word:
                continue
            if ai_word.lower() in ("ai version", "ai_version"):
                continue
            for variant in ai_word.split('/'):
                variant = variant.strip()
                if variant:
                    replacements[variant.lower()] = pr_word
    print(f"[dict] Loaded {len(replacements)} replacement pairs.")
    return replacements


# ──────────────────────────────────────────────────────────────
# 2. Run XML builders
# ──────────────────────────────────────────────────────────────

def _base_rPr(source_run_xml=None):
    """Deep-copy the source rPr (or make a blank one)."""
    if source_run_xml is not None:
        src = source_run_xml.find(qn('w:rPr'))
        if src is not None:
            return copy.deepcopy(src)
    return OxmlElement('w:rPr')


def _make_t(text: str):
    t = OxmlElement('w:t')
    t.text = text
    if text.startswith(' ') or text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return t


def build_red_run(text: str, source_run_xml=None):
    """Red colour for the original (wrong) word."""
    r = OxmlElement('w:r')
    rPr = _base_rPr(source_run_xml)

    # Remove conflicting props
    for tag in ('w:color', 'w:strike', 'w:highlight'):
        for el in rPr.findall(qn(tag)):
            rPr.remove(el)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'FF0000')
    rPr.append(color)


    r.append(rPr)
    r.append(_make_t(text))
    return r


def build_yellow_highlight_run(text: str, source_run_xml=None):
    """Yellow highlight for the corrected (PlanetRead) word."""
    r = OxmlElement('w:r')
    rPr = _base_rPr(source_run_xml)

    for tag in ('w:color', 'w:strike', 'w:highlight'):
        for el in rPr.findall(qn(tag)):
            rPr.remove(el)

    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'yellow')
    rPr.append(highlight)

    r.append(rPr)
    r.append(_make_t(text))
    return r


def build_normal_run(text: str, source_run_xml=None):
    """Plain run inheriting source formatting."""
    r = OxmlElement('w:r')
    rPr = _base_rPr(source_run_xml)
    r.append(rPr)
    r.append(_make_t(text))
    return r


# ──────────────────────────────────────────────────────────────
# 3. Per-paragraph replacement
# ──────────────────────────────────────────────────────────────

def replace_in_paragraph(para, replacements: dict):
    full_text = "".join(run.text for run in para.runs)
    if not full_text.strip():
        return 0, []

    sorted_keys = sorted(replacements.keys(), key=len, reverse=True)

    spans = []
    for key in sorted_keys:
        pattern = r'(?<![A-Za-z])' + re.escape(key) + r'(?![A-Za-z])'
        for m in re.finditer(pattern, full_text, flags=re.IGNORECASE):
            overlap = any(s < m.end() and m.start() < e for s, e, _ in spans)
            if not overlap:
                spans.append((m.start(), m.end(), replacements[key]))

    if not spans:
        return 0, []

    spans.sort(key=lambda x: x[0])
    changed = [(full_text[s:e], r) for s, e, r in spans]

    first_run_xml = para.runs[0]._r if para.runs else None
    p_xml = para._p
    for r in p_xml.findall(qn('w:r')):
        p_xml.remove(r)

    pos = 0
    for (start, end, replacement) in spans:
        # Normal text before this match
        if pos < start:
            p_xml.append(build_normal_run(full_text[pos:start], first_run_xml))

        # 1) Original word — red font
        original_word = full_text[start:end]
        p_xml.append(build_red_run(original_word, first_run_xml))

        # Space between the two words
        p_xml.append(build_normal_run(' ', first_run_xml))

        # 2) Corrected word — yellow highlight
        p_xml.append(build_yellow_highlight_run(replacement, first_run_xml))

        pos = end

    # Remaining text
    if pos < len(full_text):
        p_xml.append(build_normal_run(full_text[pos:], first_run_xml))

    return len(spans), changed


# ──────────────────────────────────────────────────────────────
# 4. Main
# ──────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) != 4:
        print("Usage: python fix_romanized.py <dictionary.docx> <input.docx> <output.docx>")
        sys.exit(1)

    dict_path, input_path, output_path = sys.argv[1], sys.argv[2], sys.argv[3]
    replacements = load_dictionary(dict_path)
    doc = Document(input_path)

    total = 0
    for para in doc.paragraphs:
        n, _ = replace_in_paragraph(para, replacements)
        total += n

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    n, _ = replace_in_paragraph(para, replacements)
                    total += n

    doc.save(output_path)
    print(f"[done] {total} replacement(s) made → {output_path}")


if __name__ == "__main__":
    main()
